import streamlit as st
import os
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
import numpy as np
from PyPDF2 import PdfReader  # Untuk membaca file PDF
from docx import Document    # Untuk membaca file DOCX

# Path ke folder 'documents'
DOCUMENTS_FOLDER = "documents"

# Daftar stop words untuk Bahasa Indonesia
STOP_WORDS_ID = [
    "dan", "di", "ke", "dari", "untuk", "yang", "pada", "dengan", "dalam", "atau", "oleh", 
    "sebagai", "adalah", "ini", "itu", "tidak", "bukan", "saya", "kami", "kita", "anda", 
    "dia", "mereka", "ada", "jika", "karena", "tetapi", "namun", "bagaimana", "mengapa",
    "apa", "dimana", "kapan", "siapa", "dapat", "harus", "akan", "sudah", "belum", "bisa"
]

# Fungsi untuk memuat daftar file dari folder 'documents'
def load_document_files(folder):
    files = []
    if os.path.exists(folder):
        files = [f for f in os.listdir(folder) if f.endswith((".txt", ".pdf", ".docx"))]
    return sorted(files)

# Fungsi untuk membaca file teks
def read_txt_file(filepath):
    with open(filepath, "r", encoding="utf-8") as file:
        return file.read()

# Fungsi untuk membaca file DOCX
def read_docx_file(filepath):
    doc = Document(filepath)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Fungsi untuk membaca file PDF
def read_pdf_file(filepath):
    pdf_reader = PdfReader(filepath)
    return "\n".join([page.extract_text() for page in pdf_reader.pages])

# Fungsi untuk membaca isi file berdasarkan ekstensi
def read_file(filepath):
    if filepath.endswith(".txt"):
        return read_txt_file(filepath)
    elif filepath.endswith(".docx"):
        return read_docx_file(filepath)
    elif filepath.endswith(".pdf"):
        return read_pdf_file(filepath)
    else:
        raise ValueError("Format file tidak didukung.")

# Fungsi untuk menyimpan teks sebagai file baru
def save_text_to_file(folder, filename, content):
    os.makedirs(folder, exist_ok=True)  # Buat folder jika belum ada
    filepath = os.path.join(folder, filename)
    with open(filepath, "w", encoding="utf-8") as file:
        file.write(content)

# Fungsi untuk menghitung TF-IDF Matrix dan skor per kalimat
def compute_tf_idf_summary(document):
    vectorizer = TfidfVectorizer(
        stop_words=STOP_WORDS_ID,  # Gunakan stop words Bahasa Indonesia
        token_pattern=r"(?u)\b\w\w+\b"  # Tokenize untuk kata dengan minimal 2 huruf
    )
    sentences = document.split('. ')  # Pisahkan dokumen menjadi kalimat
    tf_idf_matrix = vectorizer.fit_transform(sentences)

    # Hitung skor untuk setiap kalimat
    sentence_scores = tf_idf_matrix.sum(axis=1).flatten().tolist()[0]
    ranked_sentences = [(score, sentence) for sentence, score in zip(sentences, sentence_scores)]

    # Urutkan kalimat berdasarkan skor TF-IDF
    ranked_sentences = sorted(ranked_sentences, key=lambda x: x[0], reverse=True)

    # Pilih beberapa kalimat teratas sebagai ringkasan
    summary_sentences = [sentence for _, sentence in ranked_sentences[:3]]  # Ambil 3 kalimat teratas
    summary = '. '.join(summary_sentences)
    return summary

# Streamlit App
st.set_page_config(
    page_title="Perangkuman Dokumen",
    page_icon="📄",
    layout="wide",
)

# Header
st.markdown(
    """
    <style>
    .main-header {
        background-color: #0d6efd;
        color: white;
        text-align: center;
        padding: 10px;
        border-radius: 10px;
    }
    .content-box {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1);
    }
    </style>
    <div class="main-header">
        <h1>📄 Perangkuman Dokumen</h1>
        <p>Kelola dan ringkas dokumen Anda dengan mudah</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# Sidebar: Pilih atau Tambah Dokumen
st.sidebar.title("📂 Daftar Dokumen")

# **Input Data Baru** dari User
st.sidebar.subheader("✍️ Tambah Dokumen Baru")

# Pilihan Input
input_option = st.sidebar.radio("Pilih Cara Input Data", ["Unggah File", "Input Teks Manual"])

## Jika User Mengunggah File
if input_option == "Unggah File":
    uploaded_file = st.sidebar.file_uploader("Unggah File (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        content = None

        # Proses file berdasarkan tipe
        if file_extension == "txt":
            # Jika file teks, baca sebagai string
            content = uploaded_file.getvalue().decode("utf-8")
        elif file_extension == "docx":
            # Jika file DOCX, baca konten menggunakan python-docx
            doc = Document(uploaded_file)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == "pdf":
            # Jika file PDF, baca konten menggunakan PyPDF2
            pdf_reader = PdfReader(uploaded_file)
            content = "\n".join([page.extract_text() for page in pdf_reader.pages])

        if content:
            # Simpan file ke folder 'documents' sebagai teks
            save_text_to_file(DOCUMENTS_FOLDER, uploaded_file.name.replace(file_extension, "txt"), content)
            st.sidebar.success(f"File '{uploaded_file.name}' berhasil diunggah dan disimpan sebagai teks!")
        else:
            st.sidebar.error("Gagal membaca isi file. Pastikan file tidak kosong atau formatnya benar.")


# Jika User Input Manual
elif input_option == "Input Teks Manual":
    new_doc_name = st.sidebar.text_input("Nama File (contoh: dokumen_baru.txt)")
    new_doc_content = st.sidebar.text_area("Isi Dokumen", height=150)
    if st.sidebar.button("Simpan Dokumen"):
        if new_doc_name and new_doc_content:
            save_text_to_file(DOCUMENTS_FOLDER, new_doc_name, new_doc_content)
            st.sidebar.success(f"File '{new_doc_name}' berhasil disimpan!")
        else:
            st.sidebar.warning("Nama file dan isi dokumen tidak boleh kosong.")

# Memuat daftar dokumen setelah input baru
document_files = load_document_files(DOCUMENTS_FOLDER)

if not document_files:
    st.sidebar.warning("Tidak ada dokumen yang ditemukan di folder 'documents'.")
else:
    # Sidebar: Pilih Dokumen yang Ada
    selected_document = st.sidebar.selectbox("Pilih Dokumen", document_files)

    # Tampilkan isi dokumen yang dipilih
    st.markdown("<div class='content-box'>", unsafe_allow_html=True)
    st.subheader(f"📄 Isi Dokumen: {selected_document}")

    document_path = os.path.join(DOCUMENTS_FOLDER, selected_document)
    document_content = read_file(document_path)

    # Menampilkan konten dokumen
    st.text_area("Konten Dokumen", document_content, height=400)

    # Tombol untuk membuat ringkasan
    if st.button("Buat Ringkasan"):
        with st.spinner("Sedang merangkum dokumen..."):
            summary = compute_tf_idf_summary(document_content)
            st.success("Ringkasan berhasil dibuat!")
            st.subheader("📃 Ringkasan Dokumen")
            st.write(summary)
    st.markdown("</div>", unsafe_allow_html=True)
