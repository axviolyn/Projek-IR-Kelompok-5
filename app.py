import streamlit as st
import os
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
import numpy as np
from PyPDF2 import PdfReader  # Untuk membaca file PDF
from docx import Document    # Untuk membaca file DOCX
from reportlab.pdfgen import canvas # Untuk membaca file PDF

# Path ke folder 'documents'
DOCUMENTS_FOLDER = "documents"

# Daftar stop words untuk Bahasa Indonesia
STOP_WORDS_ID = [
    "dan", "di", "ke", "dari", "untuk", "yang", "pada", "dengan", "dalam", "atau", "oleh", 
    "sebagai", "adalah", "ini", "itu", "tidak", "bukan", "saya", "kami", "kita", "anda", 
    "dia", "mereka", "ada", "jika", "karena", "tetapi", "namun", "bagaimana", "mengapa",
    "apa", "dimana", "kapan", "siapa", "dapat", "harus", "akan", "sudah", "belum", "bisa"
]

# Fungsi untuk memuat daftar file dari folder 'documents'
def load_document_files(folder):
    files = []
    if os.path.exists(folder):
        files = [f for f in os.listdir(folder) if f.endswith((".txt", ".pdf", ".docx"))]
    return sorted(files)

# Fungsi untuk membaca file teks
def read_txt_file(filepath):
    with open(filepath, "r", encoding="utf-8") as file:
        return file.read()

# Fungsi untuk membaca file DOCX
def read_docx_file(filepath):
    doc = Document(filepath)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Fungsi untuk membaca file PDF
def read_pdf_file(filepath):
    pdf_reader = PdfReader(filepath)
    return "\n".join([page.extract_text() for page in pdf_reader.pages])

# Fungsi untuk membaca isi file berdasarkan ekstensi
def read_file(filepath):
    if filepath.endswith(".txt"):
        return read_txt_file(filepath)
    elif filepath.endswith(".docx"):
        return read_docx_file(filepath)
    elif filepath.endswith(".pdf"):
        return read_pdf_file(filepath)
    else:
        raise ValueError("Format file tidak didukung.")

# Fungsi untuk menyimpan teks sebagai file .txt
def save_text_to_txt(folder, filename, content):
    os.makedirs(folder, exist_ok=True)  # Buat folder jika belum ada
    if not filename.endswith(".txt"):
        filename += ".txt"  # Pastikan file memiliki ekstensi .txt
    filepath = os.path.join(folder, filename)
    with open(filepath, "w", encoding="utf-8") as file:
        file.write(content)
    return filepath

# Fungsi untuk menyimpan teks sebagai file .docx
def save_text_to_docx(folder, filename, content):
    os.makedirs(folder, exist_ok=True)  # Buat folder jika belum ada
    if not filename.endswith(".docx"):
        filename += ".docx"  # Pastikan file memiliki ekstensi .docx
    filepath = os.path.join(folder, filename)
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filepath)
    return filepath

# Fungsi untuk menyimpan teks sebagai file .pdf
def save_text_to_pdf(folder, filename, content):
    os.makedirs(folder, exist_ok=True)  # Buat folder jika belum ada
    if not filename.endswith(".pdf"):
        filename += ".pdf"  # Pastikan file memiliki ekstensi .pdf
    filepath = os.path.join(folder, filename)
    c = canvas.Canvas(filepath)
    y = 800  # Posisi awal teks pada halaman PDF
    for line in content.split("\n"):
        c.drawString(100, y, line)
        y -= 15  # Jarak antar baris
        if y < 50:  # Jika melebihi halaman, tambahkan halaman baru
            c.showPage()
            y = 800
    c.save()
    return filepath

# Fungsi utama untuk menyimpan dokumen berdasarkan format pilihan
def save_document(folder, filename, content, format_choice):
    if format_choice == "txt":
        return save_text_to_txt(folder, filename, content)
    elif format_choice == "docx":
        return save_text_to_docx(folder, filename, content)
    elif format_choice == "pdf":
        return save_text_to_pdf(folder, filename, content)
    else:
        raise ValueError("Format tidak didukung. Pilih salah satu dari 'txt', 'docx', atau 'pdf'.")


# Fungsi untuk menghitung TF-IDF Matrix dan skor per kalimat
def compute_tf_idf_summary(document):
    vectorizer = TfidfVectorizer(
        stop_words=STOP_WORDS_ID,  # Gunakan stop words Bahasa Indonesia
        token_pattern=r"(?u)\b\w\w+\b"  # Tokenize untuk kata dengan minimal 2 huruf
    )
    sentences = document.split('. ')  # Pisahkan dokumen menjadi kalimat
    tf_idf_matrix = vectorizer.fit_transform(sentences)

    # Hitung skor untuk setiap kalimat
    sentence_scores = tf_idf_matrix.sum(axis=1).flatten().tolist()[0]
    ranked_sentences = [(score, sentence) for sentence, score in zip(sentences, sentence_scores)]

    # Urutkan kalimat berdasarkan skor TF-IDF
    ranked_sentences = sorted(ranked_sentences, key=lambda x: x[0], reverse=True)

    # Pilih beberapa kalimat teratas sebagai ringkasan
    summary_sentences = [sentence for _, sentence in ranked_sentences[:3]]  # Ambil 3 kalimat teratas
    summary = '. '.join(summary_sentences)
    return summary

# Streamlit App
st.set_page_config(
    page_title="Perangkuman",
    page_icon="📄",
    layout="wide",
)

# Header
st.markdown(
    """
    <style>
    .main-header {
        background-color: #0d6efd;
        color: white;
        text-align: center;
        padding: 10px;
        border-radius: 10px;
    }
    .content-box {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1);
    }

    .st.button {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1);
    }
    </style>
    <div class="main-header">
        <h1>📄 Perangkuman </h1>
        <p>Kelola dan ringkas dokumen Anda dengan mudah</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# Sidebar: Pilih atau Tambah Dokumen
st.sidebar.title("Perangkuman")

# *Input Data Baru* dari User
st.sidebar.subheader("✍ Tambah Dokumen Baru")

# Pilihan Input
input_option = st.sidebar.radio("Pilih Cara Input Data", ["Unggah File", "Input Teks Manual"])

## Jika User Mengunggah File
if input_option == "Unggah File":
    uploaded_file = st.sidebar.file_uploader("Unggah File (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])
    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        content = None

        # Proses file berdasarkan tipe
        if file_extension == "txt":
            # Jika file teks, baca sebagai string
            content = uploaded_file.getvalue().decode("utf-8")
        elif file_extension == "docx":
            # Jika file DOCX, baca konten menggunakan python-docx
            doc = Document(uploaded_file)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == "pdf":
            # Jika file PDF, baca konten menggunakan PyPDF2
            pdf_reader = PdfReader(uploaded_file)
            content = "\n".join([page.extract_text() for page in pdf_reader.pages])

        if content:
            # Simpan file sesuai jenis aslinya
            if file_extension == "txt":
                save_text_to_txt(DOCUMENTS_FOLDER, uploaded_file.name, content)
                st.sidebar.success(f"File '{uploaded_file.name}' berhasil diunggah dan disimpan sebagai teks!")
            elif file_extension == "docx":
                save_text_to_docx(DOCUMENTS_FOLDER, uploaded_file.name, content)
                st.sidebar.success(f"File '{uploaded_file.name}' berhasil diunggah dan disimpan sebagai dokumen Word!")
            elif file_extension == "pdf":
                save_text_to_pdf(DOCUMENTS_FOLDER, uploaded_file.name, content)
                st.sidebar.success(f"File '{uploaded_file.name}' berhasil diunggah dan disimpan sebagai PDF!")
            else:
                st.sidebar.error("Format file tidak dikenali dan tidak dapat disimpan.")
        else:
            st.sidebar.error("Gagal membaca isi file. Pastikan file tidak kosong atau formatnya benar.")




# Jika User Input Manual
elif input_option == "Input Teks Manual":
    new_doc_name = st.sidebar.text_input("Nama File (contoh: dokumen_baru)")
    new_doc_content = st.sidebar.text_area("Isi Dokumen", height=150)

    # Pilihan format file
    format_choice = st.sidebar.selectbox("Pilih Format Penyimpanan", ["txt", "docx", "pdf"])

    if st.sidebar.button("Simpan Dokumen"):
        if new_doc_name and new_doc_content:
            # Tambahkan ekstensi berdasarkan format yang dipilih
            new_doc_name = f"{new_doc_name}.{format_choice}"

            # Simpan file sesuai format
            if format_choice == "txt":
                save_text_to_txt(DOCUMENTS_FOLDER, new_doc_name, new_doc_content)
            elif format_choice == "docx":
                save_text_to_docx(DOCUMENTS_FOLDER, new_doc_name, new_doc_content)
            elif format_choice == "pdf":
                save_text_to_pdf(DOCUMENTS_FOLDER, new_doc_name, new_doc_content)

            # Tampilkan pesan sukses
            st.sidebar.success(f"File '{new_doc_name}' berhasil disimpan sebagai {format_choice.upper()}!")
        else:
            st.sidebar.warning("Nama file dan isi dokumen tidak boleh kosong.")


# Memuat daftar dokumen setelah input baru
document_files = load_document_files(DOCUMENTS_FOLDER)

if not document_files:
    st.sidebar.warning("Tidak ada dokumen yang ditemukan di folder 'documents'.")
else:
    # Sidebar: Pilih Dokumen yang Ada
    selected_document = st.sidebar.selectbox("Pilih Dokumen", document_files)

    # Tampilkan isi dokumen yang dipilih
    st.markdown("<div class='content-box'>", unsafe_allow_html=True)


    document_path = os.path.join(DOCUMENTS_FOLDER, selected_document)
    document_content = read_file(document_path)

# Buat dua kolom
col1, col2 = st.columns(2)

# Kolom pertama: Konten dokumen
with col1:
    st.subheader("📄 Konten Dokumen")
    document_content = st.text_area("Isi Dokumen", document_content, height=400)
    # Tambahkan tombol di bawah teks area di kolom pertama
    if st.button("Buat Ringkasan", key="button_konten"):
        with st.spinner("Sedang merangkum dokumen..."):
            summary = compute_tf_idf_summary(document_content)

# Kolom kedua: Hasil ringkasan
with col2:
    st.subheader("📃 Ringkasan Dokumen")
    # Gunakan text_area untuk menampilkan hasil ringkasan
    if "summary" in locals():
        st.text_area("Hasil Ringkasan", summary, height=400)
    else:
        st.text_area("Hasil Ringkasan", "Belum ada ringkasan yang dibuat.", height=400)